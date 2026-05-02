"""DOCX parser using python-docx."""

from __future__ import annotations

from pathlib import Path

from rag_builder.parsers import ParsedDocument


class DocxParser:
    """Extract text from Word documents."""

    def parse(self, file_path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(file_path))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        prefix = "#" * int(level)
                    except ValueError:
                        prefix = "##"
                    paragraphs.append(f"{prefix} {text}")
                else:
                    paragraphs.append(text)

        # Extract tables as markdown
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                # Add header separator after first row
                if len(rows) > 1:
                    sep = "| " + " | ".join(["---"] * len(table.columns)) + " |"
                    rows.insert(1, sep)
                paragraphs.append("\n".join(rows))

        return ParsedDocument(
            source=str(file_path),
            content="\n\n".join(paragraphs),
            metadata={
                "file_size": file_path.stat().st_size,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
            file_type="docx",
        )
